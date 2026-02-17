"""
AI-Powered Resume Enhancement Service
Generates MS Word standard formatted resumes with AI optimization
"""

class AIResumeEnhancer:
    """Enhances resume content using AI techniques"""
    
    # Industry-standard resume formats (MS Word compatible)
    RESUME_FORMATS = {
        'professional': {
            'name': 'Professional',
            'description': 'Formal, traditional resume layout',
            'font': 'Calibri',
            'size': 11,
            'spacing': 1.15,
            'margins': {'top': 0.5, 'bottom': 0.5, 'left': 0.75, 'right': 0.75},
            'colors': {'primary': (0, 0, 0), 'accent': (0, 51, 102)},
            'sections_order': ['contact', 'summary', 'education', 'skills', 'projects', 'experience', 'personal_details']
        },
        'modern': {
            'name': 'Modern',
            'description': 'Contemporary design with visual elements',
            'font': 'Segoe UI',
            'size': 11,
            'spacing': 1.15,
            'margins': {'top': 0.75, 'bottom': 0.75, 'left': 1, 'right': 1},
            'colors': {'primary': (31, 78, 121), 'accent': (192, 0, 0)},
            'sections_order': ['contact', 'summary', 'education', 'skills', 'projects', 'experience', 'personal_details']
        },
        'simple': {
            'name': 'Simple',
            'description': 'Clean, minimalist ATS-optimized format',
            'font': 'Arial',
            'size': 10,
            'spacing': 1.0,
            'margins': {'top': 0.5, 'bottom': 0.5, 'left': 0.75, 'right': 0.75},
            'colors': {'primary': (0, 0, 0), 'accent': (64, 64, 64)},
            'sections_order': ['contact', 'summary', 'education', 'skills', 'projects', 'experience', 'personal_details']
        },
        'technical': {
            'name': 'Technical',
            'description': 'Skills-focused for tech professionals',
            'font': 'Consolas',
            'size': 10,
            'spacing': 1.0,
            'margins': {'top': 0.75, 'bottom': 0.75, 'left': 0.75, 'right': 0.75},
            'colors': {'primary': (0, 0, 0), 'accent': (0, 102, 204)},
            'sections_order': ['contact', 'summary', 'education', 'skills', 'projects', 'experience', 'personal_details']
        },
        'academic': {
            'name': 'Academic',
            'description': 'CV-style for academic purposes',
            'font': 'Times New Roman',
            'size': 12,
            'spacing': 1.5,
            'margins': {'top': 1, 'bottom': 1, 'left': 1, 'right': 1},
            'colors': {'primary': (0, 0, 0), 'accent': (51, 51, 51)},
            'sections_order': ['contact', 'summary', 'education', 'skills', 'projects', 'experience', 'personal_details']
        },
        'detailed': {
            'name': 'Detailed',
            'description': 'Comprehensive layout with all sections',
            'font': 'Garamond',
            'size': 11,
            'spacing': 1.15,
            'margins': {'top': 0.75, 'bottom': 0.75, 'left': 0.75, 'right': 0.75},
            'colors': {'primary': (0, 0, 0), 'accent': (102, 102, 102)},
            'sections_order': ['contact', 'summary', 'education', 'skills', 'projects', 'experience', 'personal_details']
        }
    }

    # Action words for impact statements
    ACTION_WORDS = {
        'leadership': ['Spearheaded', 'Orchestrated', 'Championed', 'Directed', 'Steered', 'Pioneered'],
        'improvement': ['Accelerated', 'Optimized', 'Enhanced', 'Improved', 'Streamlined', 'Refined'],
        'creation': ['Developed', 'Created', 'Designed', 'Built', 'Established', 'Launched'],
        'achievement': ['Delivered', 'Achieved', 'Accomplished', 'Exceeded', 'Excelled', 'Succeeded'],
        'analysis': ['Analyzed', 'Assessed', 'Evaluated', 'Investigated', 'Examined', 'Reviewed']
    }

    @staticmethod
    def enhance_headline(headline, industry='technology'):
        """AI-enhance job headline to be more impactful"""
        if not headline:
            return "Dedicated Professional"
        
        # Extract first word
        first_word = headline.split()[0] if headline else "Professional"
        
        enhancements = {
            'technology': f"Senior {first_word} | Full-Stack Developer",
            'business': f"Strategic {headline} | Business Solutions Expert",
            'creative': f"Creative {headline} | Design Specialist",
            'finance': f"Financial {headline} | Analytics Expert",
            'healthcare': f"Healthcare {headline} | Patient Care Focused",
            'management': f"Results-Driven {headline} | Team Leader"
        }
        
        return enhancements.get(industry, f"Experienced {headline}")

    @staticmethod
    def enhance_summary(summary, industry='technology'):
        """AI-enhance summary with action words and metrics"""
        if not summary:
            return "Dedicated professional with proven track record of driving measurable results."
        
        # Extract years of experience if mentioned
        years_mention = ""
        if 'year' in summary.lower():
            years_mention = " with extensive industry experience"
        
        # Add context based on industry
        industry_context = {
            'technology': "| Tech-savvy professional driving digital innovation",
            'business': "| Business strategist with proven ROI track record",
            'creative': "| Creative innovator with strong portfolio",
            'finance': "| Financial expert with regulatory compliance expertise",
            'healthcare': "| Patient-centered professional with clinical excellence",
            'management': "| Visionary leader building high-performing teams"
        }
        
        enhanced = summary + years_mention + ". " + industry_context.get(industry, "")
        return enhanced.strip()

    @staticmethod
    def enhance_bullet(text):
        """Convert text to professional bullet point"""
        if not text:
            return ""
        
        text = text.strip()
        
        # Remove existing bullet characters
        if text.startswith('•') or text.startswith('-') or text.startswith('*'):
            text = text.lstrip('•-* ').strip()
        
        # Capitalize first letter
        if text and not text[0].isupper():
            text = text[0].upper() + text[1:]
        
        return f"• {text}"

    @staticmethod
    def enhance_experience_bullets(experience_text):
        """Convert experience text to professional bullet points with metrics"""
        if not experience_text:
            return ""
        
        lines = experience_text.split('\n')
        enhanced_bullets = []
        
        for line in lines:
            line = line.strip()
            if line:
                enhanced = AIResumeEnhancer.enhance_bullet(line)
                
                # Add impact metric suggestions if missing
                if 'led' in line.lower() or 'managed' in line.lower():
                    if '%' not in line and 'team' not in line.lower():
                        enhanced += " [Added measurable impact metrics]"
                
                enhanced_bullets.append(enhanced)
        
        return '\n'.join(enhanced_bullets)

    @staticmethod
    def enhance_project_bullets(projects_text):
        """
        Convert projects text to professional format with separated titles and descriptions.
        
        Format: Project titles are not bulleted (bold in PDF), descriptions are bulleted.
        Expects input like:
        - "Project Title - Description text"
        - "Project Title\nDescription text"
        
        Returns formatted text with __TITLE__ marking for titles (used by PDF generator)
        """
        if not projects_text:
            return ""
        
        lines = projects_text.split('\n')
        formatted_projects = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Remove existing bullet characters
            clean_line = line.lstrip('•-* ').strip()
            
            # Check if this is a title or description
            # Heuristic: if it contains " - " it's likely "Title - Description"
            if ' - ' in clean_line:
                parts = clean_line.split(' - ', 1)
                title = parts[0].strip()
                description = parts[1].strip()
                
                # Add title without bullet (marked with special indicator for PDF)
                formatted_projects.append(f"__TITLE__{title}")
                
                # Capitalize and add bullet to description
                if description:
                    if description and not description[0].isupper():
                        description = description[0].upper() + description[1:]
                    formatted_projects.append(f"• {description}")
            else:
                # If no " - " separator, treat as project title if it's short
                # (likely a project name)
                if len(clean_line) < 80 and not clean_line.startswith('developed') and \
                   not clean_line.startswith('built') and not clean_line.startswith('created'):
                    # Looks like a title
                    formatted_projects.append(f"__TITLE__{clean_line}")
                else:
                    # Looks like a description - add bullet
                    if clean_line and not clean_line[0].isupper():
                        clean_line = clean_line[0].upper() + clean_line[1:]
                    formatted_projects.append(f"• {clean_line}")
        
        return '\n'.join(formatted_projects)

    @staticmethod
    def categorize_skills(skills_str):
        """Merge all skills into single Technical Skills list, removing duplicates
        
        Combines all user-provided skills (both Technical and Additional) into one unified
        Technical Skills category. Removes duplicates while preserving first occurrence order.
        
        Args:
            skills_str: Comma-separated skills string
            
        Returns:
            dict: Dictionary with single key 'Technical Skills' containing merged list
        """
        if not skills_str:
            return {}
        
        skills_list = [s.strip() for s in skills_str.split(',')]
        
        # Remove duplicates while preserving order (case-insensitive)
        seen = set()
        merged_skills = []
        for skill in skills_list:
            skill_lower = skill.lower()
            if skill_lower not in seen:
                seen.add(skill_lower)
                merged_skills.append(skill)
        
        # Return single category with all merged skills - NO separate Additional Skills section
        return {'Technical Skills': merged_skills}

    @staticmethod
    def format_resume_ms_word_standard(template, user, profile):
        """Generate resume in MS Word standard format"""
        format_spec = AIResumeEnhancer.RESUME_FORMATS.get(template, {})
        
        if not format_spec:
            return "Error: Invalid template"
        
        resume_parts = []
        section_divider = '─' * 70
        
        # Header - ATS-Friendly Contact Information (Clean, Minimal)
        # Build contact line with phone and email
        contact_parts = [profile.phone if profile and profile.phone else '[Your Phone]']
        contact_parts.append(profile.email if profile and profile.email else user.email)
        
        # Add only labels for social profiles (no URLs)
        if profile and profile.linkedin:
            contact_parts.append('LinkedIn')
        if profile and profile.github:
            contact_parts.append('GitHub')
        if profile and profile.leetcode:
            contact_parts.append('LeetCode')
        
        contact_line = ' | '.join(contact_parts)
        headline = profile.headline if profile and profile.headline else 'Professional Developer'
        
        resume_parts.append(f"""{user.name.upper()}
{headline}
{contact_line}""")
        
        # Add other links if provided
        if profile and profile.other_links and profile.other_links.strip():
            links = [link.strip() for link in profile.other_links.split('\n') if link.strip()]
            if links:
                resume_parts.append(f"Additional Links: {' | '.join(links)}")
        
        # Professional Summary
        if profile.summary:
            enhanced_summary = AIResumeEnhancer.enhance_summary(profile.summary)
            resume_parts.append(f"""
PROFESSIONAL SUMMARY
{section_divider}
{enhanced_summary}""")
        
        # Education Section
        if profile.education:
            resume_parts.append(f"""
EDUCATION
{section_divider}
{profile.education}""")
        
        # Skills Section
        if profile.skills:
            merged_skills = AIResumeEnhancer.categorize_skills(profile.skills)
            if merged_skills and 'Technical Skills' in merged_skills:
                resume_parts.append(f"""
SKILLS
{section_divider}""")
                skills_list = merged_skills['Technical Skills']
                resume_parts.append(f"__SKILL_LABEL__Technical Skills:__/SKILL_LABEL__ {', '.join(skills_list)}")
        
        # Projects Section
        if profile.projects:
            enhanced_proj = AIResumeEnhancer.enhance_project_bullets(profile.projects)
            resume_parts.append(f"""
PROJECTS & ACHIEVEMENTS
{section_divider}
{enhanced_proj}""")

        # Experience Section
        if profile.experience:
            enhanced_exp = AIResumeEnhancer.enhance_experience_bullets(profile.experience)
            resume_parts.append(f"""
PROFESSIONAL EXPERIENCE
{section_divider}
{enhanced_exp}""")
        
        # Personal Details Section
        personal_details = []
        if hasattr(profile, 'dob') and profile.dob:
            dob_val = profile.dob
            try:
                from datetime import datetime
                dob_date = datetime.strptime(dob_val, '%Y-%m-%d')
                dob_val = dob_date.strftime('%d %B %Y')
            except:
                pass
            personal_details.append(f"Date of Birth: {dob_val}")
        
        if hasattr(profile, 'languages') and profile.languages:
            personal_details.append(f"Languages: {profile.languages}")
            
        if hasattr(profile, 'hobbies') and profile.hobbies:
            personal_details.append(f"Hobbies: {profile.hobbies}")

        if personal_details:
            resume_parts.append(f"\nPERSONAL DETAILS\n{section_divider}\n" + "\n".join(personal_details))
        
        return '\n'.join(resume_parts)

    @staticmethod
    def generate_word_document(template, user, profile, output_path):
        """Generate MS Word (.docx) file with proper formatting"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            format_spec = AIResumeEnhancer.RESUME_FORMATS.get(template, {})
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(format_spec['margins']['top'])
                section.bottom_margin = Inches(format_spec['margins']['bottom'])
                section.left_margin = Inches(format_spec['margins']['left'])
                section.right_margin = Inches(format_spec['margins']['right'])
            
            # Add header with name (ATS-friendly: no fancy formatting)
            header = doc.add_paragraph()
            header_run = header.add_run(user.name.upper())
            header_run.font.size = Pt(14)
            header_run.font.bold = True
            # No color - keeps ATS readable
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add role tagline (1 line, no special characters)
            tagline = doc.add_paragraph("Senior Software Developer | Full-Stack Engineer")
            tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tagline_run = tagline.runs[0]
            tagline_run.font.size = Pt(10)
            
            # Add contact info (minimal formatting, pipe-separated)
            contact = doc.add_paragraph(f"{user.email} | [Your Phone] | linkedin.com/in/yourprofile | github.com/yourprofile")
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact.runs[0]
            contact_run.font.size = Pt(10)
            
            # Add professional summary
            if profile.summary:
                doc.add_heading('PROFESSIONAL SUMMARY', level=1)
                summary_para = doc.add_paragraph(AIResumeEnhancer.enhance_summary(profile.summary))
                summary_para_format = summary_para.paragraph_format
                summary_para_format.space_after = Pt(6)
            
            # Add education
            if profile.education:
                doc.add_heading('EDUCATION', level=1)
                doc.add_paragraph(profile.education)
            
            # Add skills
            if profile.skills:
                doc.add_heading('SKILLS', level=1)
                merged_skills = AIResumeEnhancer.categorize_skills(profile.skills)
                if merged_skills and 'Technical Skills' in merged_skills:
                    p = doc.add_paragraph()
                    # Add category label in bold
                    label_run = p.add_run('Technical Skills:')
                    label_run.font.bold = True
                    label_run.font.color.rgb = RGBColor(0, 0, 0)  # Dark black
                    # Add skills in normal font
                    p.add_run(f" {', '.join(merged_skills['Technical Skills'])}")
            
            # Add projects
            if profile.projects:
                doc.add_heading('PROJECTS & ACHIEVEMENTS', level=1)
                enhanced_projects = AIResumeEnhancer.enhance_project_bullets(profile.projects)
                proj_lines = enhanced_projects.split('\n')
                for line in proj_lines:
                    line = line.strip()
                    if not line:
                        continue
                    # Check if this is a project title (marked with __TITLE__)
                    if '__TITLE__' in line:
                        # Extract title and remove marker
                        title_text = line.replace('__TITLE__', '').strip()
                        # Add as bold paragraph without bullet
                        p = doc.add_paragraph(title_text)
                        p_format = p.paragraph_format
                        p_format.line_spacing = 1.15
                        for run in p.runs:
                            run.font.bold = True
                    else:
                        # Add as bulleted item
                        doc.add_paragraph(line, style='List Bullet')

            # Add experience
            if profile.experience:
                doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
                exp_lines = profile.experience.split('\n')
                for line in exp_lines:
                    doc.add_paragraph(line.strip(), style='List Bullet')
                    
            # Add personal details
            personal_details = []
            if hasattr(profile, 'dob') and profile.dob:
                dob_val = profile.dob
                try:
                    from datetime import datetime
                    dob_date = datetime.strptime(dob_val, '%Y-%m-%d')
                    dob_val = dob_date.strftime('%d %B %Y')
                except:
                    pass
                personal_details.append(f"Date of Birth: {dob_val}")
            
            if hasattr(profile, 'languages') and profile.languages:
                personal_details.append(f"Languages: {profile.languages}")
                
            if hasattr(profile, 'hobbies') and profile.hobbies:
                personal_details.append(f"Hobbies: {profile.hobbies}")

            if personal_details:
                doc.add_heading('PERSONAL DETAILS', level=1)
                for detail in personal_details:
                    doc.add_paragraph(detail)
            
            # Save document
            doc.save(output_path)
            return True
            
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return False

    @staticmethod
    def get_format_specs(template):
        """Get detailed format specifications"""
        return AIResumeEnhancer.RESUME_FORMATS.get(template, {})

    @staticmethod
    def list_all_formats():
        """List all available resume formats"""
        return [
            {
                'id': key,
                'name': value['name'],
                'description': value['description'],
                'font': value['font'],
                'size': value['size']
            }
            for key, value in AIResumeEnhancer.RESUME_FORMATS.items()
        ]
